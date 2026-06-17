import re
from datetime import date
from pathlib import Path
from shutil import copyfileobj

import fitz
import pdfplumber
from docx import Document
from fastapi import UploadFile
from openpyxl import load_workbook
from sqlalchemy import delete, or_
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models import ReferenceDocument, ReferenceDocumentChunk


SUPPORTED_TYPES = {".pdf", ".docx", ".xlsx"}
PARSING_COMPLETED = "Parsed"
PARSING_REVIEW_REQUIRED = "Parsing Review Required"
INDEXED = "Indexed"
INDEX_REVIEW_REQUIRED = "Indexing Review Required"


def save_reference_document(
    db: Session,
    upload: UploadFile,
    title: str | None = None,
    category: str = "Other",
    effective_date: date | None = None,
    version_label: str | None = None,
    source_type: str | None = None,
    notes: str | None = None,
    uploaded_by: str = "system",
    parse_immediately: bool = True,
) -> ReferenceDocument:
    suffix = Path(upload.filename or "").suffix.lower()
    if suffix not in SUPPORTED_TYPES:
        raise ValueError("Supported reference formats are PDF, DOCX, and XLSX.")

    settings = get_settings()
    library_root = Path(settings.upload_dir) / "reference-library"
    library_root.mkdir(parents=True, exist_ok=True)
    target = _unique_path(library_root / (upload.filename or f"reference{suffix}"))
    with target.open("wb") as handle:
        copyfileobj(upload.file, handle)

    document = ReferenceDocument(
        title=(title or target.stem).strip(),
        category=category or "Other",
        file_name=upload.filename or target.name,
        file_path=str(target),
        file_type=suffix,
        effective_date=effective_date,
        version_label=version_label,
        source_type=source_type or "Uploaded Reference",
        uploaded_by=uploaded_by or "system",
        parsing_status="Pending",
        indexed_status="Pending",
        notes=notes,
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    if parse_immediately:
        parse_reference_document(db, document.id)
        db.refresh(document)
    return document


def parse_reference_document(db: Session, document_id: int) -> ReferenceDocument:
    document = db.get(ReferenceDocument, document_id)
    if not document:
        raise ValueError("Reference document not found")

    db.execute(delete(ReferenceDocumentChunk).where(ReferenceDocumentChunk.reference_document_id == document_id))
    db.flush()
    try:
        pages = _extract_pages(Path(document.file_path), document.file_type)
        chunks = chunk_reference_text(document_id, pages)
        for item in chunks:
            db.add(ReferenceDocumentChunk(**item))
        document.parsing_status = PARSING_COMPLETED if pages else PARSING_REVIEW_REQUIRED
        document.indexed_status = INDEXED if chunks else INDEX_REVIEW_REQUIRED
    except Exception as exc:
        document.parsing_status = PARSING_REVIEW_REQUIRED
        document.indexed_status = INDEX_REVIEW_REQUIRED
        document.notes = _append_note(document.notes, f"Possible review required during parsing: {exc}")
    db.commit()
    db.refresh(document)
    return document


def extract_pdf_text(file_path: str | Path) -> list[dict]:
    try:
        return _extract_pdf_text_with_pymupdf(file_path)
    except Exception:
        return _extract_pdf_text_with_pdfplumber(file_path)


def _extract_pdf_text_with_pymupdf(file_path: str | Path) -> list[dict]:
    pages = []
    with fitz.open(str(file_path)) as pdf:
        for index, page in enumerate(pdf, start=1):
            text = (page.get_text("text") or "").strip()
            pages.append({"page_number": index, "text": text})
    return pages


def _extract_pdf_text_with_pdfplumber(file_path: str | Path) -> list[dict]:
    pages = []
    with pdfplumber.open(str(file_path)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            pages.append({"page_number": index, "text": text})
    return pages


def chunk_reference_text(document_id: int, pages: list[dict]) -> list[dict]:
    chunks = []
    chunk_index = 1
    for page in pages:
        page_number = page.get("page_number")
        text = _normalise_text(page.get("text") or "")
        if not text:
            continue
        for part in _split_text(text):
            rule_number, section_number, heading = _identify_reference(part)
            chunks.append({
                "reference_document_id": document_id,
                "page_number": page_number,
                "section_number": section_number,
                "rule_number": rule_number,
                "heading": heading,
                "content_text": part,
                "chunk_index": chunk_index,
            })
            chunk_index += 1
    return chunks


def search_reference_library(db: Session, query: str, category: str | None = None) -> list[dict]:
    clean_query = (query or "").strip()
    if not clean_query:
        return []
    like = f"%{clean_query}%"
    db_query = db.query(ReferenceDocumentChunk, ReferenceDocument).join(ReferenceDocument)
    db_query = db_query.filter(or_(
        ReferenceDocumentChunk.content_text.ilike(like),
        ReferenceDocumentChunk.heading.ilike(like),
        ReferenceDocumentChunk.rule_number.ilike(like),
        ReferenceDocumentChunk.section_number.ilike(like),
        ReferenceDocument.title.ilike(like),
    ))
    if category:
        db_query = db_query.filter(ReferenceDocument.category == category)
    results = []
    for chunk, document in db_query.order_by(ReferenceDocument.title.asc(), ReferenceDocumentChunk.page_number.asc(), ReferenceDocumentChunk.chunk_index.asc()).limit(100).all():
        results.append({
            "document_id": document.id,
            "document_title": document.title,
            "category": document.category,
            "page_number": chunk.page_number,
            "rule_number": chunk.rule_number,
            "section_number": chunk.section_number,
            "heading": chunk.heading,
            "matching_text_snippet": _snippet(chunk.content_text, clean_query),
        })
    return results


def get_reference_document(db: Session, document_id: int) -> ReferenceDocument:
    document = db.get(ReferenceDocument, document_id)
    if not document:
        raise ValueError("Reference document not found")
    return document


def get_reference_document_chunks(db: Session, document_id: int) -> list[ReferenceDocumentChunk]:
    return db.query(ReferenceDocumentChunk).filter(ReferenceDocumentChunk.reference_document_id == document_id).order_by(ReferenceDocumentChunk.page_number.asc(), ReferenceDocumentChunk.chunk_index.asc()).all()


def suggested_reference_matches(db: Session, exception_type: str) -> list[dict]:
    return search_reference_library(db, exception_type or "")[:5]


def _extract_pages(path: Path, file_type: str) -> list[dict]:
    if file_type == ".pdf":
        return extract_pdf_text(path)
    if file_type == ".docx":
        doc = Document(str(path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return [{"page_number": 1, "text": text}]
    if file_type == ".xlsx":
        wb = load_workbook(path, read_only=True, data_only=True)
        pages = []
        for sheet_index, sheet in enumerate(wb.worksheets, start=1):
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value not in (None, "")]
                if values:
                    rows.append(" | ".join(values))
            pages.append({"page_number": sheet_index, "text": "\n".join(rows)})
        return pages
    return []


def _identify_reference(text: str) -> tuple[str | None, str | None, str | None]:
    first_line = next((line.strip() for line in text.splitlines() if line.strip()), "")
    rule_match = re.search(r"\bRule\s+([0-9A-Za-z.-]+)\b[:.\-\s]*(.*)", first_line, re.IGNORECASE)
    section_match = re.search(r"\bSection\s+([0-9A-Za-z().-]+)\b[:.\-\s]*(.*)", first_line, re.IGNORECASE)
    numbered_rule = re.match(r"^([0-9]{1,3}[A-Za-z]?)\.\s+(.+)", first_line)
    if rule_match:
        return rule_match.group(1), None, _trim_heading(rule_match.group(2) or first_line)
    if section_match:
        return None, section_match.group(1), _trim_heading(section_match.group(2) or first_line)
    if numbered_rule:
        return numbered_rule.group(1), None, _trim_heading(numbered_rule.group(2))
    return None, None, _trim_heading(first_line)


def _split_text(text: str, size: int = 1800) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\n{2,}", text) if part.strip()]
    chunks = []
    current = ""
    for paragraph in paragraphs or [text]:
        if len(current) + len(paragraph) + 2 <= size:
            current = f"{current}\n\n{paragraph}".strip()
            continue
        if current:
            chunks.append(current)
        current = paragraph
        while len(current) > size:
            chunks.append(current[:size].strip())
            current = current[size:].strip()
    if current:
        chunks.append(current)
    return chunks


def _snippet(text: str, query: str, width: int = 320) -> str:
    lower = text.lower()
    needle = query.lower()
    index = lower.find(needle)
    if index < 0:
        return text[:width]
    start = max(index - 90, 0)
    end = min(index + len(query) + width - 90, len(text))
    prefix = "..." if start else ""
    suffix = "..." if end < len(text) else ""
    return f"{prefix}{text[start:end].strip()}{suffix}"


def _normalise_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _trim_heading(value: str | None) -> str | None:
    text = _normalise_text(value or "")
    if not text:
        return None
    return text[:500]


def _append_note(existing: str | None, note: str) -> str:
    if existing:
        return f"{existing}\n{note}"
    return note


def _unique_path(path: Path) -> Path:
    target = path
    counter = 1
    while target.exists():
        target = path.with_name(f"{path.stem}-{counter}{path.suffix}")
        counter += 1
    return target
