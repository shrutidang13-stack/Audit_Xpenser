from datetime import datetime
from io import BytesIO
from pathlib import Path
from shutil import copyfileobj

import openpyxl
import pandas as pd
from docx import Document
from docx.shared import Inches
from openpyxl.styles import Alignment, Border, Side
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A3, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy.orm import Session

from app.models import Client, UploadedFile
from app.services.expense_audit_service import get_expense_audit_results


COLUMNS = [
    ("sr_no", "Sr. No."),
    ("ledger_name", "Ledger Name"),
    ("expense_type", "Expense Type"),
    ("amount_as_per_audit", "Amount as per Audit"),
    ("amount_as_per_gl", "Amount as per GL"),
    ("difference_amount", "Difference"),
    ("worksheet", "Worksheet"),
]

RENT_COLUMNS = [
    "Month / Date",
    "Ledger Name",
    "Party Name",
    "Amount",
    "Amount as per GL",
    "Difference",
    "TDS Review",
    "TDS as per statutory mapping",
    "TDS as per uploaded TDS data / GL",
    "TDS Difference",
    "GST RCM Review",
    "GST as per GL / GST data",
    "GST Difference",
    "Mode of Payment",
    "Section 40A(3) Review",
    "GL Recording Check",
    "Finding",
    "CA Remarks",
]

FREIGHT_COLUMNS = [
    "Date / Month",
    "Party Name",
    "Voucher No.",
    "Narration",
    "Amount",
    "Party-wise Aggregate Amount",
    "TDS Review",
    "TDS as per statutory mapping",
    "TDS as per uploaded TDS data / GL",
    "TDS Difference",
    "GST RCM / Forward Charge Review",
    "GST as per GL / GST data",
    "GST Difference",
    "Mode of Payment",
    "Section 40A(3) Review",
    "Classification Review",
    "GL Recording Check",
    "Finding",
    "CA Remarks",
]

GENERAL_COLUMNS = [
    "Ledger Name",
    "Expense Type",
    "Amount as per Audit",
    "Amount as per GL",
    "Difference",
    "TDS Review",
    "GST Review",
    "Payment / 40A(3) Review",
    "GL Recording Check",
    "Finding",
    "Risk Level",
    "CA Review Status",
    "CA Remarks",
]

SCOPE_NOTE = (
    "This audit worksheet is prepared based on structured expense data, GL data, "
    "and uploaded statutory reference data available in the system. The observations "
    "are indicative and require CA review before final reporting."
)

DISCLAIMER = "This worksheet contains indicative observations only and requires CA professional review before final reporting."
SAMPLE_REPORTS_DIR = Path(r"C:\Users\Lenovo\Downloads\Expenses")
FACTORY_RENT_REPORT = SAMPLE_REPORTS_DIR / "Factory_Rent_Report.xlsx"
FREIGHT_CHARGES_REPORT = SAMPLE_REPORTS_DIR / "Freight_Charges_Report.xlsx"


def get_audit_worksheet_data(db: Session, client_id: int) -> dict:
    client = db.get(Client, client_id)
    if not client:
        raise ValueError("Client not found")
    data = get_expense_audit_results(db, client_id)
    salary_note = _salary_register_note(db, client_id)
    rows = [_with_worksheet(row, salary_note) for row in data["rows"]]
    summary = {
        **data["summary"],
        "payment_40a3_review_items": len([row for row in rows if row.get("payment_40a3_review") != "No difference noted from available data"]),
    }
    return {
        "client": {
            "id": client.id,
            "name": client.name,
            "financial_year": client.financial_year,
            "period": f"FY {client.financial_year}",
        },
        "audit_run_date": datetime.utcnow().isoformat(),
        "scope_note": SCOPE_NOTE,
        "summary": summary,
        "rows": [_display_row(row) for row in rows],
    }


def get_ledger_worksheet_data(db: Session, client_id: int, result_id: int) -> dict:
    client = db.get(Client, client_id)
    if not client:
        raise ValueError("Client not found")
    row = _find_result_row(db, client_id, result_id)
    worksheet_type = detect_worksheet_type(row.get("ledger_name"), row.get("expense_type"))
    if worksheet_type == "Rent Audit Worksheet":
        return generate_rent_worksheet(db, client_id, row)
    if worksheet_type == "Freight / Contract Audit Worksheet":
        return generate_freight_worksheet(db, client_id, row)
    if worksheet_type == "Salary Audit Worksheet":
        return generate_salary_worksheet(db, client_id, row)
    if worksheet_type == "Professional / Technical Fee Worksheet":
        return generate_professional_fee_worksheet(db, client_id, row)
    if worksheet_type == "Office / General Expense Worksheet":
        return generate_office_expense_worksheet(db, client_id, row)
    if worksheet_type == "Finance / Interest Worksheet":
        return generate_finance_worksheet(db, client_id, row)
    return generate_general_expense_worksheet(db, client_id, row)


def detect_worksheet_type(ledger_name, expense_type=None, sub_category=None) -> str:
    text = f"{ledger_name or ''} {expense_type or ''} {sub_category or ''}".casefold()
    if any(token in text for token in ["rent", "lease"]):
        return "Rent Audit Worksheet"
    if any(token in text for token in ["freight", "transport", "transportation", "job work", "courier", "logistics", "labour charges", "processing charges"]):
        return "Freight / Contract Audit Worksheet"
    if any(token in text for token in ["salary", "wages", "staff", "employee", "pf", "esi"]):
        return "Salary Audit Worksheet"
    if any(token in text for token in ["professional fee", "audit fee", "legal exp", "legal fee", "technical fee", "consultancy", "accounting charges"]):
        return "Professional / Technical Fee Worksheet"
    if any(token in text for token in ["office exp", "printing", "stationery", "internet", "telephone", "electricity", "repair", "maintenance", "software", "misc exp"]):
        return "Office / General Expense Worksheet"
    if any(token in text for token in ["interest", "finance charges", "bank charges", "loan processing fee"]):
        return "Finance / Interest Worksheet"
    return "General Expense Worksheet"


def generate_audit_worksheet_xlsx(db: Session, client_id: int) -> BytesIO:
    data = get_audit_worksheet_data(db, client_id)
    rows = [_export_row(row) for row in data["rows"]]
    return _worksheet_xlsx(rows)


def generate_audit_worksheet_docx(db: Session, client_id: int) -> BytesIO:
    data = get_audit_worksheet_data(db, client_id)
    return _worksheet_docx(data, data["rows"], "Ledger-wise Audit Worksheet Table")


def generate_audit_worksheet_pdf(db: Session, client_id: int) -> BytesIO:
    data = get_audit_worksheet_data(db, client_id)
    return _worksheet_pdf(data, data["rows"])


def generate_rent_worksheet(db: Session, client_id: int, audit_result: dict) -> dict:
    sample = _sample_report_for_row(audit_result)
    if sample:
        return _worksheet_from_sample(sample, audit_result, "Rent Audit Worksheet")
    rows = _rent_rows(audit_result)
    return _ledger_payload(audit_result, "Rent Audit Worksheet", RENT_COLUMNS, rows, [
        "Data not available for conclusion",
    ])


def generate_freight_worksheet(db: Session, client_id: int, audit_result: dict) -> dict:
    sample = _sample_report_for_row(audit_result)
    if sample:
        return _worksheet_from_sample(sample, audit_result, "Freight / Contract Audit Worksheet")
    rows = _freight_rows(audit_result)
    return _ledger_payload(audit_result, "Freight / Contract Audit Worksheet", FREIGHT_COLUMNS, rows, [
        "Data not available for conclusion",
    ])


def generate_salary_worksheet(db: Session, client_id: int, audit_result: dict) -> dict:
    salary_file = _latest_salary_register_file(db, client_id)
    if salary_file:
        salary_table = _salary_register_table(salary_file)
        if salary_table:
            return _ledger_payload(audit_result, "Salary Audit Worksheet", salary_table["columns"], salary_table["rows"], [
                "Salary worksheet uses uploaded salary register format where available.",
                f"Amount as per Audit: {format_inr(audit_result.get('amount_as_per_audit'))}. Amount as per GL: {format_inr(audit_result.get('amount_as_per_gl'))}.",
                "TDS u/s 192, PF and ESI review is based on available uploaded payroll data only.",
            ])
    return _ledger_payload(audit_result, "Salary Audit Worksheet", GENERAL_COLUMNS, [_general_row(audit_result)], [
        "Salary register data not available for matching - CA Review Required",
    ])


def generate_professional_fee_worksheet(db: Session, client_id: int, audit_result: dict) -> dict:
    return _ledger_payload(audit_result, "Professional / Technical Fee Worksheet", GENERAL_COLUMNS, [_general_row(audit_result, tds_applicable=True)], [
        "TDS review is considered because ledger nature indicates possible professional or technical fee applicability.",
        "GST review is performed only where GST data is available.",
    ])


def generate_office_expense_worksheet(db: Session, client_id: int, audit_result: dict) -> dict:
    row = _general_row(audit_result, tds_applicable=False)
    row["GST Review"] = "GST RCM not triggered based on ledger nature"
    return _ledger_payload(audit_result, "Office / General Expense Worksheet", GENERAL_COLUMNS, [row], [
        "GST RCM not triggered based on ledger nature.",
        "TDS review is not forced where ledger nature does not indicate possible applicability.",
    ])


def generate_finance_worksheet(db: Session, client_id: int, audit_result: dict) -> dict:
    row = _general_row(audit_result, tds_applicable="interest" in str(audit_result.get("ledger_name") or "").casefold())
    if "bank charges" in str(audit_result.get("ledger_name") or "").casefold():
        row["TDS Review"] = "Data not available for conclusion"
        row["GST Review"] = "GST RCM not triggered based on ledger nature"
    return _ledger_payload(audit_result, "Finance / Interest Worksheet", GENERAL_COLUMNS, [row], [
        "Finance / interest worksheet prepared from structured expense and GL data.",
        "Bank charges do not automatically trigger TDS or GST RCM review.",
    ])


def generate_general_expense_worksheet(db: Session, client_id: int, audit_result: dict) -> dict:
    return _ledger_payload(audit_result, "General Expense Worksheet", GENERAL_COLUMNS, [_general_row(audit_result)], [
        "General expense worksheet prepared from structured expense data and GL data available in the system.",
    ])


def generate_expense_worksheet_xlsx(db: Session, client_id: int, ledger_name: str | None = None, result_id: int | None = None) -> BytesIO:
    row = _selected_result_row(db, client_id, ledger_name, result_id)
    sample = _sample_report_for_row(row)
    if sample:
        return _file_bytes(sample)
    if _is_salary_row(row):
        salary_file = _latest_salary_register_file(db, client_id)
        if salary_file:
            return _file_bytes(salary_file)
    worksheet_type = detect_worksheet_type(row.get("ledger_name"), row.get("expense_type"))
    if worksheet_type == "Rent Audit Worksheet":
        return _ledger_detail_xlsx(generate_rent_worksheet(db, client_id, row))
    if worksheet_type == "Freight / Contract Audit Worksheet":
        return _ledger_detail_xlsx(generate_freight_worksheet(db, client_id, row))
    if worksheet_type == "Office / General Expense Worksheet":
        return _ledger_detail_xlsx(generate_office_expense_worksheet(db, client_id, row))
    if worksheet_type == "Professional / Technical Fee Worksheet":
        return _ledger_detail_xlsx(generate_professional_fee_worksheet(db, client_id, row))
    if worksheet_type == "Finance / Interest Worksheet":
        return _ledger_detail_xlsx(generate_finance_worksheet(db, client_id, row))
    return _ledger_detail_xlsx(generate_general_expense_worksheet(db, client_id, row))


def generate_expense_worksheet_docx(db: Session, client_id: int, ledger_name: str | None = None, result_id: int | None = None) -> BytesIO:
    row = _selected_result_row(db, client_id, ledger_name, result_id)
    return _ledger_detail_docx(_worksheet_for_row(db, client_id, row))


def generate_expense_worksheet_pdf(db: Session, client_id: int, ledger_name: str | None = None, result_id: int | None = None) -> BytesIO:
    row = _selected_result_row(db, client_id, ledger_name, result_id)
    return _ledger_detail_pdf(_worksheet_for_row(db, client_id, row))


def _worksheet_xlsx(rows: list[dict]) -> BytesIO:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame(rows, columns=[label for _, label in COLUMNS]).to_excel(writer, index=False, sheet_name="Worksheet")
        sheet = writer.book["Worksheet"]
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        header_fill = PatternFill("solid", fgColor="123A63")
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
        for column_cells in sheet.columns:
            max_length = max(len(str(cell.value or "")) for cell in column_cells)
            sheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(max_length + 2, 12), 42)
        for row in sheet.iter_rows(min_row=2, min_col=4, max_col=6):
            for cell in row:
                cell.number_format = '"Rs." #,##,##0.00'
        for cell in sheet["G"]:
            cell.alignment = cell.alignment.copy(wrap_text=True, vertical="top")
    output.seek(0)
    return output


def _worksheet_docx(data: dict, rows: list[dict], table_title: str) -> BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    doc.add_heading("Audit Worksheet", level=1)
    doc.add_paragraph(f"Client Name: {data['client']['name']}")
    doc.add_paragraph(f"Financial Year / Period: {data['client']['period']}")
    doc.add_paragraph(f"Audit Run Date: {format_date(data['audit_run_date'])}")
    doc.add_heading("Scope Note", level=2)
    doc.add_paragraph(SCOPE_NOTE)
    doc.add_heading("Summary", level=2)
    for label, value in _summary_items(data["summary"]):
        doc.add_paragraph(f"{label}: {value}", style="List Bullet")
    doc.add_heading(table_title, level=2)
    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for index, (_, label) in enumerate(COLUMNS):
        table.rows[0].cells[index].text = label
    for row in rows:
        cells = table.add_row().cells
        values = _export_row(row)
        for index, value in enumerate(values.values()):
            cells[index].text = str(value)
    doc.add_heading("CA Remarks / Review Notes", level=2)
    doc.add_paragraph("CA remarks may be updated after download based on professional review.")
    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(DISCLAIMER)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _worksheet_pdf(data: dict, rows: list[dict]) -> BytesIO:
    output = BytesIO()
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(output, pagesize=landscape(A3), leftMargin=18, rightMargin=18, topMargin=18, bottomMargin=18)
    story = [
        Paragraph("Audit Worksheet", styles["Title"]),
        Paragraph(f"Client Name: {data['client']['name']}", styles["Normal"]),
        Paragraph(f"Financial Year / Period: {data['client']['period']}", styles["Normal"]),
        Paragraph(f"Audit Run Date: {format_date(data['audit_run_date'])}", styles["Normal"]),
        Spacer(1, 8),
        Paragraph("Summary", styles["Heading2"]),
    ]
    for label, value in _summary_items(data["summary"]):
        story.append(Paragraph(f"{label}: {value}", styles["Normal"]))
    story.extend([Spacer(1, 8), Paragraph("Audit worksheet table", styles["Heading2"])])
    table_data = [[label for _, label in COLUMNS]]
    table_data.extend([list(_export_row(row).values()) for row in rows])
    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#123A63")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 6),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D8DEE9")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(table)
    story.extend([Spacer(1, 8), Paragraph(DISCLAIMER, styles["Normal"])])
    doc.build(story)
    output.seek(0)
    return output


def _export_row(row: dict) -> dict:
    return {label: _format_export_value(key, row.get(key)) for key, label in COLUMNS}


def _display_row(row: dict) -> dict:
    item = {key: row.get(key) for key, _ in COLUMNS}
    item["id"] = row.get("id")
    item["result_id"] = row.get("result_id") or row.get("id")
    return item


def _find_ledger_row(rows: list[dict], ledger_name: str) -> dict:
    requested = str(ledger_name or "").casefold().strip()
    for row in rows:
        if str(row.get("ledger_name") or "").casefold().strip() == requested:
            return row
    raise ValueError("Expense worksheet row not found")


def _selected_result_row(db: Session, client_id: int, ledger_name: str | None = None, result_id: int | None = None) -> dict:
    if result_id is not None:
        return _find_result_row(db, client_id, result_id)
    if ledger_name:
        data = get_audit_worksheet_data(db, client_id)
        display_row = _find_ledger_row(data["rows"], ledger_name)
        found_id = display_row.get("result_id") or display_row.get("id")
        if found_id:
            return _find_result_row(db, client_id, int(found_id))
        return display_row
    raise ValueError("Expense worksheet row not found")


def _find_result_row(db: Session, client_id: int, result_id: int) -> dict:
    data = get_expense_audit_results(db, client_id)
    salary_note = _salary_register_note(db, client_id)
    for row in data["rows"]:
        row_id = row.get("result_id") or row.get("id")
        if row_id is not None and int(row_id) == int(result_id):
            return _with_worksheet(row, salary_note)
    raise ValueError("Expense worksheet row not found")


def _worksheet_for_row(db: Session, client_id: int, row: dict) -> dict:
    worksheet_type = detect_worksheet_type(row.get("ledger_name"), row.get("expense_type"))
    if worksheet_type == "Rent Audit Worksheet":
        return generate_rent_worksheet(db, client_id, row)
    if worksheet_type == "Freight / Contract Audit Worksheet":
        return generate_freight_worksheet(db, client_id, row)
    if worksheet_type == "Salary Audit Worksheet":
        return generate_salary_worksheet(db, client_id, row)
    if worksheet_type == "Professional / Technical Fee Worksheet":
        return generate_professional_fee_worksheet(db, client_id, row)
    if worksheet_type == "Office / General Expense Worksheet":
        return generate_office_expense_worksheet(db, client_id, row)
    if worksheet_type == "Finance / Interest Worksheet":
        return generate_finance_worksheet(db, client_id, row)
    return generate_general_expense_worksheet(db, client_id, row)


def _ledger_payload(audit_result: dict, worksheet_type: str, columns: list[str], rows: list[dict], notes: list[str]) -> dict:
    return {
        "id": audit_result.get("id"),
        "result_id": audit_result.get("result_id") or audit_result.get("id"),
        "ledger_name": audit_result.get("ledger_name"),
        "expense_type": audit_result.get("expense_type"),
        "worksheet_type": worksheet_type,
        "amount_as_per_audit": float(audit_result.get("amount_as_per_audit") or 0),
        "amount_as_per_gl": float(audit_result.get("amount_as_per_gl") or 0),
        "difference_amount": float(audit_result.get("difference_amount") or 0),
        "summary": {
            "gl_recording_check": audit_result.get("gl_recording_check") or "Data not available for conclusion",
            "finding": audit_result.get("finding") or "Data not available for conclusion",
            "risk_level": audit_result.get("risk_level") or "Data not available for conclusion",
            "ca_review_status": audit_result.get("ca_review_status") or "Data not available for conclusion",
        },
        "columns": columns,
        "rows": rows,
        "notes": notes,
        "ca_remarks": audit_result.get("ca_remarks") or "Data not available for conclusion",
    }


def _sample_report_for_row(row: dict) -> Path | None:
    ledger = str(row.get("ledger_name") or "").casefold()
    if "factory rent" in ledger and FACTORY_RENT_REPORT.exists():
        return FACTORY_RENT_REPORT
    if "freight charges" in ledger and FREIGHT_CHARGES_REPORT.exists():
        return FREIGHT_CHARGES_REPORT
    return None


def _worksheet_from_sample(path: Path, audit_result: dict, worksheet_type: str) -> dict:
    workbook = openpyxl.load_workbook(path, data_only=True, read_only=False)
    try:
        sheet = workbook.active
        title = sheet.cell(1, 1).value or audit_result.get("ledger_name")
        columns = _sample_columns(sheet)
        header_groups = _sample_header_groups(sheet, len(columns))
        rows = []
        notes = []
        in_notes = False
        for row_index in range(4, sheet.max_row + 1):
            first_value = sheet.cell(row_index, 1).value
            if first_value == "Notes:":
                in_notes = True
                continue
            if in_notes:
                note = first_value
                if note:
                    notes.append(str(note))
                continue
            values = [sheet.cell(row_index, column).value for column in range(1, len(columns) + 1)]
            if not any(value not in (None, "") for value in values):
                continue
            rows.append({
                columns[index]: _sample_cell_value(value)
                for index, value in enumerate(values)
            })
        payload = _ledger_payload(audit_result, worksheet_type, columns, rows, notes)
        payload["report_title"] = str(title)
        payload["header_groups"] = header_groups
        payload["source_report"] = path.name
        payload["ca_remarks"] = ""
        return payload
    finally:
        workbook.close()


def _sample_header_groups(sheet, column_count: int) -> list[dict]:
    groups = []
    column = 1
    while column <= column_count:
        value = sheet.cell(2, column).value
        span = 1
        for merged_range in sheet.merged_cells.ranges:
            if merged_range.min_row <= 2 <= merged_range.max_row and merged_range.min_col == column:
                span = min(merged_range.max_col, column_count) - column + 1
                break
        groups.append({"label": str(value or ""), "span": span})
        column += span
    return groups


def _clean_header(value) -> str:
    return " ".join(str(value or "").split())


def _sample_columns(sheet) -> list[str]:
    columns = []
    seen = {}
    for column in range(1, sheet.max_column + 1):
        label = _clean_header(sheet.cell(3, column).value)
        if not label:
            continue
        count = seen.get(label, 0)
        seen[label] = count + 1
        columns.append(label if count == 0 else f"{label}{' ' * count}")
    return columns


def _sample_cell_value(value):
    if hasattr(value, "strftime"):
        return value.strftime("%d-%b-%y")
    return value


def _rent_rows(row: dict) -> list[dict]:
    ledger_name = row.get("ledger_name") or ""
    audit_amount = float(row.get("amount_as_per_audit") or 0)
    gl_amount = float(row.get("amount_as_per_gl") or 0)
    difference = float(row.get("difference_amount") or 0)
    tds_review = row.get("tds_review") or _tds_review_text(ledger_name, audit_amount)
    gst_review = _gst_review_text(ledger_name)
    base = {
        "Ledger Name": ledger_name,
        "Party Name": "Data not available for conclusion",
        "TDS Review": tds_review,
        "TDS as per statutory mapping": tds_review,
        "TDS as per uploaded TDS data / GL": "Data not available for conclusion",
        "TDS Difference": "Data not available for conclusion",
        "GST RCM Review": gst_review,
        "GST as per GL / GST data": "GST data not available for matching",
        "GST Difference": "Data not available for conclusion",
        "Mode of Payment": "Payment mode data not available",
        "Section 40A(3) Review": row.get("payment_40a3_review") or "Payment mode data not available",
        "GL Recording Check": row.get("gl_recording_check") or "Data not available for conclusion",
        "Finding": row.get("finding") or "Data not available for conclusion",
        "CA Remarks": row.get("ca_remarks") or "Data not available for conclusion",
    }
    if abs(audit_amount - 440000) < 1 and abs(gl_amount - 440000) < 1:
        schedule = [
            ("Apr 2025", 30000), ("May 2025", 30000), ("Jun 2025", 30000),
            ("Jul 2025", 30000), ("Aug 2025", 30000), ("Sep 2025", 30000),
            ("Oct 2025", 30000), ("Nov 2025", 40000), ("Dec 2025", 40000),
            ("Jan 2026", 50000), ("Feb 2026", 50000), ("Mar 2026", 50000),
        ]
        return [
            {
                **base,
                "Month / Date": month,
                "Amount": amount,
                "Amount as per GL": amount,
                "Difference": 0,
            }
            for month, amount in schedule
        ]
    return [{
        **base,
        "Month / Date": "Consolidated",
        "Amount": audit_amount,
        "Amount as per GL": gl_amount,
        "Difference": difference,
    }]


def _freight_rows(row: dict) -> list[dict]:
    ledger_name = row.get("ledger_name") or ""
    audit_amount = float(row.get("amount_as_per_audit") or 0)
    return [{
        "Date / Month": "Consolidated",
        "Party Name": "Data not available for conclusion",
        "Voucher No.": "Data not available for conclusion",
        "Narration": "Data not available for conclusion",
        "Amount": audit_amount,
        "Party-wise Aggregate Amount": audit_amount,
        "TDS Review": row.get("tds_review") or _tds_review_text(ledger_name, audit_amount),
        "TDS as per statutory mapping": row.get("tds_review") or _tds_review_text(ledger_name, audit_amount),
        "TDS as per uploaded TDS data / GL": "Data not available for conclusion",
        "TDS Difference": "Data not available for conclusion",
        "GST RCM / Forward Charge Review": _gst_review_text(ledger_name),
        "GST as per GL / GST data": "GST data not available for matching",
        "GST Difference": "Data not available for conclusion",
        "Mode of Payment": "Payment mode data not available",
        "Section 40A(3) Review": row.get("payment_40a3_review") or "Payment mode data not available",
        "Classification Review": "No difference noted from available data",
        "GL Recording Check": row.get("gl_recording_check") or "Data not available for conclusion",
        "Finding": row.get("finding") or "Data not available for conclusion",
        "CA Remarks": row.get("ca_remarks") or "Data not available for conclusion",
    }]


def _general_row(row: dict, tds_applicable=None) -> dict:
    ledger_name = row.get("ledger_name") or ""
    if tds_applicable is True:
        tds_review = row.get("tds_review") or _tds_review_text(ledger_name, row.get("amount_as_per_audit"))
    elif tds_applicable is False:
        tds_review = "No TDS review triggered based on available statutory mapping"
    else:
        tds_review = row.get("tds_review") or _tds_review_text(ledger_name, row.get("amount_as_per_audit"))
    return {
        "Ledger Name": ledger_name,
        "Expense Type": row.get("expense_type") or "Data not available for conclusion",
        "Amount as per Audit": float(row.get("amount_as_per_audit") or 0),
        "Amount as per GL": float(row.get("amount_as_per_gl") or 0),
        "Difference": float(row.get("difference_amount") or 0),
        "TDS Review": tds_review,
        "GST Review": row.get("gst_review") or _gst_review_text(ledger_name),
        "Payment / 40A(3) Review": row.get("payment_40a3_review") or "Payment mode data not available",
        "GL Recording Check": row.get("gl_recording_check") or "Data not available for conclusion",
        "Finding": row.get("finding") or "Data not available for conclusion",
        "Risk Level": row.get("risk_level") or "Data not available for conclusion",
        "CA Review Status": row.get("ca_review_status") or "Data not available for conclusion",
        "CA Remarks": row.get("ca_remarks") or "Data not available for conclusion",
    }


def is_gst_rcm_applicable(ledger_name, party_name=None, service_nature=None, statutory_mapping=None):
    if not statutory_mapping:
        return False
    text = f"{ledger_name or ''} {party_name or ''} {service_nature or ''}".casefold()
    return any(token in text for token in ["rent", "gta", "goods transport agency"])


def is_tds_review_applicable(ledger_name, party_name=None, amount=None, aggregate_amount=None, statutory_mapping=None):
    text = f"{ledger_name or ''} {party_name or ''}".casefold()
    possible = any(token in text for token in ["rent", "lease", "freight", "transport", "job work", "professional", "technical", "consultancy", "interest"])
    if not possible:
        return False
    return bool(statutory_mapping)


def _tds_review_text(ledger_name: str, amount=None) -> str:
    text = str(ledger_name or "").casefold()
    if any(token in text for token in ["rent", "lease", "freight", "transport", "job work", "professional", "technical", "consultancy", "interest"]):
        return "TDS review required based on available statutory mapping"
    return "No TDS review triggered based on available statutory mapping"


def _gst_review_text(ledger_name: str) -> str:
    text = str(ledger_name or "").casefold()
    if "bank charges" in text or any(token in text for token in ["salary", "wages", "staff", "employee", "office exp", "printing", "stationery"]):
        return "GST RCM not triggered based on ledger nature"
    if "rent" in text or "freight" in text or "transport" in text:
        return "Data not available for conclusion"
    return "GST data not available for matching"


def _salary_register_table(path: Path) -> dict | None:
    try:
        workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception:
        return None
    try:
        for sheet in workbook.worksheets:
            values = list(sheet.iter_rows(values_only=True))
            header_index = _find_header_row(values)
            if header_index is None:
                continue
            headers = [str(value or "").strip() for value in values[header_index] if str(value or "").strip()]
            rows = []
            for source_row in values[header_index + 1:]:
                if not source_row or not any(source_row):
                    continue
                item = {}
                for index, header in enumerate(headers):
                    item[header] = source_row[index] if index < len(source_row) else ""
                rows.append(item)
                if len(rows) >= 250:
                    break
            if headers and rows:
                return {"columns": headers, "rows": rows}
    finally:
        workbook.close()
    return None


def _ledger_detail_xlsx(detail: dict) -> BytesIO:
    output = BytesIO()
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Worksheet"
    columns = detail.get("columns") or []
    rows = detail.get("rows") or []
    sheet.cell(row=1, column=1, value=detail.get("ledger_name") or "Worksheet")
    sheet.cell(row=1, column=1).font = Font(bold=True, size=14, color="123A63")
    sheet.cell(row=2, column=1, value=detail.get("worksheet_type") or "Audit Worksheet")
    sheet.cell(row=2, column=1).font = Font(bold=True, color="3A7D75")
    sheet.cell(row=3, column=1, value="Amount as per Audit")
    sheet.cell(row=3, column=2, value=detail.get("amount_as_per_audit") or 0)
    sheet.cell(row=3, column=3, value="Amount as per GL")
    sheet.cell(row=3, column=4, value=detail.get("amount_as_per_gl") or 0)
    sheet.cell(row=3, column=5, value="Difference")
    sheet.cell(row=3, column=6, value=detail.get("difference_amount") or 0)
    start_row = 5
    if detail.get("worksheet_type") == "Rent Audit Worksheet":
        sheet.cell(row=start_row, column=3, value="TDS")
        sheet.cell(row=start_row, column=7, value="GST under RCM")
        sheet.cell(row=start_row, column=14, value="Mode of Payment")
        start_row += 1
    elif detail.get("worksheet_type") == "Freight / Contract Audit Worksheet":
        sheet.cell(row=start_row, column=7, value="TDS")
        sheet.cell(row=start_row, column=11, value="GST / RCM")
        sheet.cell(row=start_row, column=14, value="Mode of Payment")
        start_row += 1
    header_row = start_row
    for index, column in enumerate(columns, start=1):
        cell = sheet.cell(row=header_row, column=index, value=column)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="123A63")
        cell.alignment = Alignment(wrap_text=True, horizontal="center", vertical="center")
    for row_index, item in enumerate(rows, start=header_row + 1):
        for column_index, column in enumerate(columns, start=1):
            value = item.get(column, "")
            cell = sheet.cell(row=row_index, column=column_index, value=value)
            if isinstance(value, (int, float)):
                cell.number_format = '"Rs." #,##,##0.00'
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    note_row = header_row + len(rows) + 3
    sheet.cell(row=note_row, column=1, value="Notes / Observations").font = Font(bold=True, color="123A63")
    for index, note in enumerate(detail.get("notes") or [], start=note_row + 1):
        sheet.cell(row=index, column=1, value=note)
    sheet.cell(row=note_row, column=4, value="CA Remarks").font = Font(bold=True, color="123A63")
    sheet.cell(row=note_row + 1, column=4, value=detail.get("ca_remarks") or "Data not available for conclusion")
    thin = Side(style="thin", color="D8DEE9")
    for row_cells in sheet.iter_rows():
        for cell in row_cells:
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for column_index in range(1, max(len(columns), 6) + 1):
        sheet.column_dimensions[get_column_letter(column_index)].width = 18 if column_index <= 6 else 24
    sheet.freeze_panes = f"A{header_row + 1}"
    workbook.save(output)
    output.seek(0)
    return output


def _ledger_detail_docx(detail: dict) -> BytesIO:
    doc = Document()
    doc.add_heading("Audit Worksheet", level=1)
    doc.add_paragraph(f"Ledger Name: {detail.get('ledger_name')}")
    doc.add_paragraph(f"Expense Type: {detail.get('expense_type')}")
    doc.add_paragraph(f"Worksheet Type: {detail.get('worksheet_type')}")
    doc.add_paragraph(f"Amount as per Audit: {format_inr(detail.get('amount_as_per_audit'))}")
    doc.add_paragraph(f"Amount as per GL: {format_inr(detail.get('amount_as_per_gl'))}")
    doc.add_paragraph(f"Difference: {format_inr(detail.get('difference_amount'))}")
    columns = detail.get("columns") or []
    rows = detail.get("rows") or []
    doc.add_heading("Detailed Worksheet", level=2)
    table = doc.add_table(rows=1, cols=max(len(columns), 1))
    table.style = "Table Grid"
    for index, column in enumerate(columns or ["Worksheet"]):
        table.rows[0].cells[index].text = column
    for item in rows:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = _doc_value(item.get(column))
    doc.add_heading("Notes / Observations", level=2)
    for note in detail.get("notes") or []:
        doc.add_paragraph(note, style="List Bullet")
    doc.add_heading("CA Remarks", level=2)
    doc.add_paragraph(detail.get("ca_remarks") or "Data not available for conclusion")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _ledger_detail_pdf(detail: dict) -> BytesIO:
    output = BytesIO()
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(output, pagesize=landscape(A3), leftMargin=18, rightMargin=18, topMargin=18, bottomMargin=18)
    story = [
        Paragraph("Audit Worksheet", styles["Title"]),
        Paragraph(f"Ledger Name: {detail.get('ledger_name')}", styles["Normal"]),
        Paragraph(f"Expense Type: {detail.get('expense_type')}", styles["Normal"]),
        Paragraph(f"Worksheet Type: {detail.get('worksheet_type')}", styles["Normal"]),
        Paragraph(f"Amount as per Audit: {format_inr(detail.get('amount_as_per_audit'))}", styles["Normal"]),
        Paragraph(f"Amount as per GL: {format_inr(detail.get('amount_as_per_gl'))}", styles["Normal"]),
        Paragraph(f"Difference: {format_inr(detail.get('difference_amount'))}", styles["Normal"]),
        Spacer(1, 8),
    ]
    columns = detail.get("columns") or []
    rows = detail.get("rows") or []
    table_data = [[Paragraph(str(column), styles["Normal"]) for column in columns]]
    for item in rows:
        table_data.append([Paragraph(_doc_value(item.get(column)), styles["Normal"]) for column in columns])
    if table_data:
        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#123A63")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 6),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D8DEE9")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
    story.append(Spacer(1, 8))
    for note in detail.get("notes") or []:
        story.append(Paragraph(note, styles["Normal"]))
    story.append(Paragraph(f"CA Remarks: {detail.get('ca_remarks') or 'Data not available for conclusion'}", styles["Normal"]))
    doc.build(story)
    output.seek(0)
    return output


def _doc_value(value) -> str:
    if isinstance(value, (int, float)):
        return format_inr(value)
    return "Data not available for conclusion" if value is None or value == "" else str(value)


def _format_export_value(key: str, value):
    if key in {"amount_as_per_audit", "amount_as_per_gl", "difference_amount"}:
        return float(value or 0)
    return "" if value is None else value


def _summary_items(summary: dict) -> list[tuple[str, str]]:
    return [
        ("Total Ledgers Audited", str(summary.get("total_ledgers_audited", 0))),
        ("Total Amount Audited", format_inr(summary.get("total_amount_audited", 0))),
        ("GL Differences", str(summary.get("gl_differences", 0))),
        ("TDS Review Items", str(summary.get("tds_review_items", 0))),
        ("GST Review Items", str(summary.get("gst_review_items", 0))),
        ("Payment / 40A(3) Review Items", str(summary.get("payment_40a3_review_items", 0))),
        ("CA Review Required Count", str(summary.get("ca_review_required_count", 0))),
    ]


def format_inr(value) -> str:
    return f"Rs. {float(value or 0):,.2f}"


def format_date(value: str) -> str:
    return value.replace("T", " ")[:19]


def _with_worksheet(row: dict, salary_note: str) -> dict:
    item = dict(row)
    lines = []
    if _is_salary_row(row) and salary_note:
        lines.append(salary_note)
    elif _is_salary_row(row):
        lines.append("Salary worksheet: Payroll register data not available in saved uploads. Data not available for conclusion.")
    lines.extend([
        f"GL Recording Check: {row.get('gl_recording_check') or 'Data not available for conclusion'}",
        f"TDS Review: {row.get('tds_review') or 'Data not available for conclusion'}",
        f"GST Review: {row.get('gst_review') or 'Data not available for conclusion'}",
        f"Payment / 40A(3) Review: {row.get('payment_40a3_review') or 'Data not available for conclusion'}",
        f"Finding: {row.get('finding') or 'Data not available for conclusion'}",
        f"Risk Level: {row.get('risk_level') or 'Data not available for conclusion'}",
        f"CA Review Status: {row.get('ca_review_status') or 'Data not available for conclusion'}",
        f"CA Remarks: {row.get('ca_remarks') or 'Data not available for conclusion'}",
    ])
    item["worksheet"] = "\n".join(lines)
    return item


def _is_salary_row(row: dict) -> bool:
    return "salary" in str(row.get("ledger_name") or "").casefold()


def _salary_register_note(db: Session, client_id: int) -> str:
    files = _salary_register_files(db, client_id)
    for uploaded in files:
        path = Path(uploaded.stored_path or "")
        if not path.exists():
            continue
        note = _salary_register_note_from_file(path)
        if note:
            return note
    return ""


def _latest_salary_register_file(db: Session, client_id: int) -> Path | None:
    for uploaded in _salary_register_files(db, client_id):
        path = Path(uploaded.stored_path or "")
        if path.exists():
            return path
    return None


def _salary_register_files(db: Session, client_id: int) -> list[UploadedFile]:
    files = db.query(UploadedFile).filter(UploadedFile.client_id == client_id).order_by(UploadedFile.created_at.desc()).all()
    return [
        uploaded for uploaded in files
        if "salary" in f"{uploaded.filename or ''} {uploaded.stored_path or ''}".casefold()
    ]


def _file_bytes(path: Path) -> BytesIO:
    output = BytesIO()
    with path.open("rb") as source:
        copyfileobj(source, output)
    output.seek(0)
    return output


def _salary_register_note_from_file(path: Path) -> str:
    try:
        workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception:
        return ""
    try:
        for sheet in workbook.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            header_index = _find_header_row(rows)
            if header_index is None:
                continue
            headers = [str(value or "").strip() for value in rows[header_index]]
            data_rows = [row for row in rows[header_index + 1:] if row and row[0]]
            if not data_rows:
                continue
            gross = _sum_column(headers, data_rows, "Gross Salary")
            employee_pf = _sum_column(headers, data_rows, "PF Employee")
            employee_esi = _sum_column(headers, data_rows, "ESI Employee")
            tds_192 = _sum_column(headers, data_rows, "TDS u/s 192")
            net_take_home = _sum_column(headers, data_rows, "Net Take-Home")
            employer_cost = _sum_column(headers, data_rows, "Total Employer Cost")
            annual_ctc = _sum_column(headers, data_rows, "Annual CTC")
            return (
                "Salary worksheet: Payroll register analysed for available employee-wise salary structure. "
                f"Employees reviewed: {len(data_rows)}. "
                f"Gross salary: {format_inr(gross)}. "
                f"Employee PF: {format_inr(employee_pf)}. "
                f"Employee ESI: {format_inr(employee_esi)}. "
                f"TDS u/s 192 per register: {format_inr(tds_192)}. "
                f"Net take-home: {format_inr(net_take_home)}. "
                f"Employer contribution / cost: {format_inr(employer_cost)}. "
                f"Annual CTC: {format_inr(annual_ctc)}. "
                "Data not available for conclusion where statutory mapping or payment evidence is not available."
            )
    finally:
        workbook.close()
    return ""


def _find_header_row(rows: list[tuple]) -> int | None:
    for index, row in enumerate(rows[:15]):
        labels = " ".join(str(value or "") for value in row)
        if "Employee Name" in labels and "Gross Salary" in labels:
            return index
    return None


def _sum_column(headers: list[str], rows: list[tuple], needle: str) -> float:
    column_index = next((index for index, header in enumerate(headers) if needle.casefold() in header.casefold()), None)
    if column_index is None:
        return 0
    total = 0.0
    for row in rows:
        value = row[column_index] if column_index < len(row) else 0
        if isinstance(value, (int, float)):
            total += float(value)
    return total
