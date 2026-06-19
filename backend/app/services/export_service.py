from io import BytesIO

import pandas as pd
from docx import Document
from sqlalchemy.orm import Session

from app.models import Client, ClientQuery, Form3CDImpact, RiskScore, StatutoryAlert, WorkingPaper
from app.services.gst_reco_service import export_rows


def client_queries_excel(db: Session, client_id: int) -> BytesIO:
    rows = [_to_dict(q, ["query_number", "ledger", "vendor", "transaction_date", "amount", "issue_detected", "required_document", "priority", "status", "suggested_wording"]) for q in db.query(ClientQuery).filter(ClientQuery.client_id == client_id).all()]
    return _excel(rows, "Client Queries")


def exception_report_excel(db: Session, client_id: int) -> BytesIO:
    scores = db.query(RiskScore).filter(RiskScore.client_id == client_id).all()
    alerts = db.query(StatutoryAlert).filter(StatutoryAlert.client_id == client_id).all()
    impacts = db.query(Form3CDImpact).filter(Form3CDImpact.client_id == client_id).all()
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame([_to_dict(s, ["transaction_id", "score", "level", "reasons"]) for s in scores]).to_excel(writer, index=False, sheet_name="Risk Scores")
        pd.DataFrame([_to_dict(a, ["transaction_id", "alert_type", "issue", "severity", "suggested_review"]) for a in alerts]).to_excel(writer, index=False, sheet_name="Statutory Alerts")
        pd.DataFrame([_to_dict(i, ["source_type", "source_id", "clause_area", "observation", "suggested_review"]) for i in impacts]).to_excel(writer, index=False, sheet_name="Form 3CD Impact")
    output.seek(0)
    return output


def gst_reco_excel(db: Session, client_id: int) -> BytesIO:
    return _excel(export_rows(db, client_id), "GST Reco")


def working_paper_docx(db: Session, client_id: int) -> BytesIO:
    client = db.get(Client, client_id)
    paper = db.query(WorkingPaper).filter(WorkingPaper.client_id == client_id).order_by(WorkingPaper.id.desc()).first()
    queries = db.query(ClientQuery).filter(ClientQuery.client_id == client_id).all()
    doc = Document()
    doc.add_heading("AuditXpenser Expense Audit Working Paper", level=1)
    doc.add_paragraph(f"Client name: {client.name if client else client_id}")
    doc.add_paragraph(f"Financial year: {client.financial_year if client else ''}")
    for heading in ["Objective", "Scope", "Data uploaded and reviewed", "Expense areas analysed", "Procedures performed", "Bill matching summary", "Key exceptions", "TDS/GST/RCM observations", "Form 3CD potential impact", "Client queries generated", "CA review notes", "Conclusion placeholder", "Prepared by", "Reviewed by"]:
        doc.add_heading(heading, level=2)
        if paper and heading in {"Objective", "Scope", "Procedures performed", "Conclusion placeholder"}:
            doc.add_paragraph(paper.content)
        elif heading == "Client queries generated":
            for query in queries:
                doc.add_paragraph(f"{query.query_number}: {query.suggested_wording}", style="List Bullet")
        else:
            doc.add_paragraph("CA Review Required.")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _excel(rows: list[dict], sheet_name: str) -> BytesIO:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame(rows).to_excel(writer, index=False, sheet_name=sheet_name)
    output.seek(0)
    return output


def _to_dict(obj, fields: list[str]) -> dict:
    return {field: getattr(obj, field) for field in fields}
