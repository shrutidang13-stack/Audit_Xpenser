import re
from datetime import date
from pathlib import Path
from shutil import copyfileobj
from urllib.request import Request, urlopen

import fitz
import pdfplumber
from docx import Document
from fastapi import UploadFile
from openpyxl import load_workbook
from sqlalchemy import delete, or_
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models import ReferenceDocument, ReferenceDocumentChunk


SUPPORTED_TYPES = {".pdf", ".docx", ".xlsx"}
PARSING_COMPLETED = "Parsed"
PARSING_REVIEW_REQUIRED = "Parsing Review Required"
INDEXED = "Indexed"
INDEX_REVIEW_REQUIRED = "Indexing Review Required"

OFFICIAL_GST_CIRCULARS = (
    {
        "number": "237/31/2024-GST",
        "date": date(2024, 10, 15),
        "subject": "Implementation of sub-sections (5) and (6) of section 16 of the CGST Act, 2017",
        "file_name": "circular-no-237-2024.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-10/circular-no-237-2024.pdf",
    },
    {
        "number": "216/10/2024-GST",
        "date": date(2024, 6, 26),
        "subject": "GST liability and ITC availability in cases involving warranty and extended warranty",
        "file_name": "circular-no-216-10-2024.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-09/circular-no-216-10-2024.pdf",
    },
    {
        "number": "211/05/2024-GST",
        "date": date(2024, 6, 26),
        "subject": "Time limit under section 16(4) for RCM supplies received from unregistered persons",
        "file_name": "circular-no-211-05-2024.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-09/circular-no-211-05-2024.pdf",
    },
    {
        "number": "198/10/2023-GST",
        "date": date(2023, 7, 17),
        "subject": "Clarification on issues pertaining to e-invoice",
        "file_name": "circular-cgst-198.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-06/circular-cgst-198.pdf",
    },
    {
        "number": "195/07/2023-GST",
        "date": date(2023, 7, 17),
        "subject": "ITC availability for warranty replacement of parts and repair services",
        "file_name": "circular-cgst-195.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-06/circular-cgst-195.pdf",
    },
    {
        "number": "193/05/2023-GST",
        "date": date(2023, 7, 17),
        "subject": "Difference between ITC availed in GSTR-3B and GSTR-2A for 01-04-2019 to 31-12-2021",
        "file_name": "circular-cgst-193.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-06/circular-cgst-193.pdf",
    },
    {
        "number": "183/15/2022-GST",
        "date": date(2022, 12, 27),
        "subject": "Difference between ITC availed in GSTR-3B and GSTR-2A for FY 2017-18 and FY 2018-19",
        "file_name": "cir-183-15-2022-cgst.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-06/cir-183-15-2022-cgst.pdf",
    },
    {
        "number": "171/03/2022-GST",
        "date": date(2022, 7, 6),
        "subject": "Demand and penalty provisions for transactions involving fake invoices",
        "file_name": "cir-171-03-2022-cgst.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-06/cir-171-03-2022-cgst.pdf",
    },
    {
        "number": "170/02/2022-GST",
        "date": date(2022, 7, 6),
        "subject": "Reporting ineligible or blocked ITC and reversals in GSTR-3B and inter-State supplies in GSTR-1",
        "file_name": "cir-170-02-2022-cgst.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-06/cir-170-02-2022-cgst.pdf",
    },
    {
        "number": "160/16/2021-GST",
        "date": date(2021, 9, 20),
        "subject": "Clarification on debit notes, physical invoice copies for e-invoices, and section 16(4)",
        "file_name": "circular-no-160-16-2021-gst.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-06/circular_20no._20160_14_2021_gst.pdf",
    },
    {
        "number": "Corrigendum to 160/16/2021-GST",
        "date": date(2021, 9, 21),
        "subject": "Corrigendum to Circular No. 160/16/2021-GST",
        "file_name": "circular-no-160-16-2021-gst-corrigendum.pdf",
        "url": "https://gstcouncil.gov.in/sites/default/files/2024-06/circular_no_160_14_2021_gst_corri.pdf",
    },
)


def sync_official_gst_circulars(db: Session) -> list[ReferenceDocument]:
    """Add only curated, software-relevant GST circulars without altering user references."""
    settings = get_settings()
    library_root = Path(settings.upload_dir) / "reference-library" / "official-gst-circulars"
    library_root.mkdir(parents=True, exist_ok=True)
    synced = []

    for item in OFFICIAL_GST_CIRCULARS:
        title = f"Circular No. {item['number']}"
        document = db.query(ReferenceDocument).filter(
            ReferenceDocument.title == title,
            ReferenceDocument.source_type == "Official GST Council Circular",
        ).one_or_none()
        target = library_root / item["file_name"]
        if not target.exists():
            _download_official_pdf(item["url"], target)

        if document is None:
            document = ReferenceDocument(
                title=title,
                category="Circular / Notification",
                file_name=item["file_name"],
                file_path=str(target),
                file_type=".pdf",
                effective_date=item["date"],
                version_label=item["date"].strftime("%d-%m-%Y"),
                source_type="Official GST Council Circular",
                uploaded_by="system",
                parsing_status="Pending",
                indexed_status="Pending",
                notes=f"{item['subject']}\nOfficial source: {item['url']}",
            )
            db.add(document)
            db.commit()
            db.refresh(document)
        else:
            document.file_path = str(target)
            document.effective_date = item["date"]
            document.version_label = item["date"].strftime("%d-%m-%Y")
            document.notes = f"{item['subject']}\nOfficial source: {item['url']}"
            db.commit()

        if document.parsing_status != PARSING_COMPLETED or not document.chunks:
            document = parse_reference_document(db, document.id)
        synced.append(document)
    return synced


def _download_official_pdf(url: str, target: Path) -> None:
    request = Request(url, headers={"User-Agent": "AuditXpenser/1.0 statutory-reference-sync"})
    with urlopen(request, timeout=45) as response:
        content = response.read()
    if not content.startswith(b"%PDF"):
        raise ValueError(f"Official GST circular did not return a PDF: {url}")
    target.write_bytes(content)


def save_reference_document(
    db: Session,
    upload: UploadFile,
    title: str | None = None,
    category: str = "Other",
    effective_date: date | None = None,
    version_label: str | None = None,
    source_type: str | None = None,
    notes: str | None = None,
    uploaded_by: str = "system",
    parse_immediately: bool = True,
) -> ReferenceDocument:
    suffix = Path(upload.filename or "").suffix.lower()
    if suffix not in SUPPORTED_TYPES:
        raise ValueError("Supported reference formats are PDF, DOCX, and XLSX.")

    settings = get_settings()
    library_root = Path(settings.upload_dir) / "reference-library"
    library_root.mkdir(parents=True, exist_ok=True)
    target = _unique_path(library_root / (upload.filename or f"reference{suffix}"))
    with target.open("wb") as handle:
        copyfileobj(upload.file, handle)

    document = ReferenceDocument(
        title=(title or target.stem).strip(),
        category=category or "Other",
        file_name=upload.filename or target.name,
        file_path=str(target),
        file_type=suffix,
        effective_date=effective_date,
        version_label=version_label,
        source_type=source_type or "Uploaded Reference",
        uploaded_by=uploaded_by or "system",
        parsing_status="Pending",
        indexed_status="Pending",
        notes=notes,
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    if parse_immediately:
        parse_reference_document(db, document.id)
        db.refresh(document)
    return document


def parse_reference_document(db: Session, document_id: int) -> ReferenceDocument:
    document = db.get(ReferenceDocument, document_id)
    if not document:
        raise ValueError("Reference document not found")

    db.execute(delete(ReferenceDocumentChunk).where(ReferenceDocumentChunk.reference_document_id == document_id))
    db.flush()
    try:
        pages = _extract_pages(Path(document.file_path), document.file_type)
        chunks = chunk_reference_text(document_id, pages)
        for item in chunks:
            db.add(ReferenceDocumentChunk(**item))
        document.parsing_status = PARSING_COMPLETED if pages else PARSING_REVIEW_REQUIRED
        document.indexed_status = INDEXED if chunks else INDEX_REVIEW_REQUIRED
    except Exception as exc:
        document.parsing_status = PARSING_REVIEW_REQUIRED
        document.indexed_status = INDEX_REVIEW_REQUIRED
        document.notes = _append_note(document.notes, f"Possible review required during parsing: {exc}")
    db.commit()
    db.refresh(document)
    return document


def extract_pdf_text(file_path: str | Path) -> list[dict]:
    try:
        return _extract_pdf_text_with_pymupdf(file_path)
    except Exception:
        return _extract_pdf_text_with_pdfplumber(file_path)


def _extract_pdf_text_with_pymupdf(file_path: str | Path) -> list[dict]:
    pages = []
    with fitz.open(str(file_path)) as pdf:
        for index, page in enumerate(pdf, start=1):
            text = (page.get_text("text") or "").strip()
            pages.append({"page_number": index, "text": text})
    return pages


def _extract_pdf_text_with_pdfplumber(file_path: str | Path) -> list[dict]:
    pages = []
    with pdfplumber.open(str(file_path)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            pages.append({"page_number": index, "text": text})
    return pages


def chunk_reference_text(document_id: int, pages: list[dict]) -> list[dict]:
    chunks = []
    chunk_index = 1
    for page in pages:
        page_number = page.get("page_number")
        text = _normalise_text(page.get("text") or "")
        if not text:
            continue
        for part in _split_text(text):
            rule_number, section_number, heading = _identify_reference(part)
            chunks.append({
                "reference_document_id": document_id,
                "page_number": page_number,
                "section_number": section_number,
                "rule_number": rule_number,
                "heading": heading,
                "content_text": part,
                "chunk_index": chunk_index,
            })
            chunk_index += 1
    return chunks


def search_reference_library(db: Session, query: str, category: str | None = None) -> list[dict]:
    clean_query = (query or "").strip()
    if not clean_query:
        return []
    like = f"%{clean_query}%"
    db_query = db.query(ReferenceDocumentChunk, ReferenceDocument).join(ReferenceDocument)
    db_query = db_query.filter(or_(
        ReferenceDocumentChunk.content_text.ilike(like),
        ReferenceDocumentChunk.heading.ilike(like),
        ReferenceDocumentChunk.rule_number.ilike(like),
        ReferenceDocumentChunk.section_number.ilike(like),
        ReferenceDocument.title.ilike(like),
    ))
    if category:
        db_query = db_query.filter(ReferenceDocument.category == category)
    results = []
    for chunk, document in db_query.order_by(ReferenceDocument.title.asc(), ReferenceDocumentChunk.page_number.asc(), ReferenceDocumentChunk.chunk_index.asc()).limit(100).all():
        results.append({
            "document_id": document.id,
            "document_title": document.title,
            "category": document.category,
            "page_number": chunk.page_number,
            "rule_number": chunk.rule_number,
            "section_number": chunk.section_number,
            "heading": chunk.heading,
            "matching_text_snippet": _snippet(chunk.content_text, clean_query),
        })
    return results


def get_reference_document(db: Session, document_id: int) -> ReferenceDocument:
    document = db.get(ReferenceDocument, document_id)
    if not document:
        raise ValueError("Reference document not found")
    return document


def get_reference_document_chunks(db: Session, document_id: int) -> list[ReferenceDocumentChunk]:
    return db.query(ReferenceDocumentChunk).filter(ReferenceDocumentChunk.reference_document_id == document_id).order_by(ReferenceDocumentChunk.page_number.asc(), ReferenceDocumentChunk.chunk_index.asc()).all()


def suggested_reference_matches(db: Session, exception_type: str) -> list[dict]:
    return search_reference_library(db, exception_type or "")[:5]


def _extract_pages(path: Path, file_type: str) -> list[dict]:
    if file_type == ".pdf":
        return extract_pdf_text(path)
    if file_type == ".docx":
        doc = Document(str(path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return [{"page_number": 1, "text": text}]
    if file_type == ".xlsx":
        wb = load_workbook(path, read_only=True, data_only=True)
        pages = []
        for sheet_index, sheet in enumerate(wb.worksheets, start=1):
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value not in (None, "")]
                if values:
                    rows.append(" | ".join(values))
            pages.append({"page_number": sheet_index, "text": "\n".join(rows)})
        return pages
    return []


def _identify_reference(text: str) -> tuple[str | None, str | None, str | None]:
    first_line = next((line.strip() for line in text.splitlines() if line.strip()), "")
    rule_match = re.search(r"\bRule\s+([0-9A-Za-z.-]+)\b[:.\-\s]*(.*)", first_line, re.IGNORECASE)
    section_match = re.search(r"\bSection\s+([0-9A-Za-z().-]+)\b[:.\-\s]*(.*)", first_line, re.IGNORECASE)
    numbered_rule = re.match(r"^([0-9]{1,3}[A-Za-z]?)\.\s+(.+)", first_line)
    if rule_match:
        return rule_match.group(1), None, _trim_heading(rule_match.group(2) or first_line)
    if section_match:
        return None, section_match.group(1), _trim_heading(section_match.group(2) or first_line)
    if numbered_rule:
        return numbered_rule.group(1), None, _trim_heading(numbered_rule.group(2))
    return None, None, _trim_heading(first_line)


def _split_text(text: str, size: int = 1800) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\n{2,}", text) if part.strip()]
    chunks = []
    current = ""
    for paragraph in paragraphs or [text]:
        if len(current) + len(paragraph) + 2 <= size:
            current = f"{current}\n\n{paragraph}".strip()
            continue
        if current:
            chunks.append(current)
        current = paragraph
        while len(current) > size:
            chunks.append(current[:size].strip())
            current = current[size:].strip()
    if current:
        chunks.append(current)
    return chunks


def _snippet(text: str, query: str, width: int = 320) -> str:
    lower = text.lower()
    needle = query.lower()
    index = lower.find(needle)
    if index < 0:
        return text[:width]
    start = max(index - 90, 0)
    end = min(index + len(query) + width - 90, len(text))
    prefix = "..." if start else ""
    suffix = "..." if end < len(text) else ""
    return f"{prefix}{text[start:end].strip()}{suffix}"


def _normalise_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _trim_heading(value: str | None) -> str | None:
    text = _normalise_text(value or "")
    if not text:
        return None
    return text[:500]


def _append_note(existing: str | None, note: str) -> str:
    if existing:
        return f"{existing}\n{note}"
    return note


def _unique_path(path: Path) -> Path:
    target = path
    counter = 1
    while target.exists():
        target = path.with_name(f"{path.stem}-{counter}{path.suffix}")
        counter += 1
    return target
