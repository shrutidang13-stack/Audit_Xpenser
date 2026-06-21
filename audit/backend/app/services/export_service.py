from io import BytesIO

import pandas as pd
from docx import Document
from sqlalchemy.orm import Session

from app.models import BillMatchingResult, Client, ClientQuery, FixedAssetDepreciation, FixedAssetReviewAlert, Form3CDImpact, RiskScore, StatutoryAlert, WorkingPaper
from app.services.gst_reco_service import export_rows


def client_queries_excel(db: Session, client_id: int) -> BytesIO:
    rows = [_to_dict(q, ["query_number", "ledger", "vendor", "transaction_date", "amount", "issue_detected", "required_document", "priority", "status", "suggested_wording"]) for q in db.query(ClientQuery).filter(ClientQuery.client_id == client_id).all()]
    return _excel(rows, "Client Queries")


def exception_report_excel(db: Session, client_id: int) -> BytesIO:
    scores = db.query(RiskScore).filter(RiskScore.client_id == client_id).all()
    alerts = db.query(StatutoryAlert).filter(StatutoryAlert.client_id == client_id).all()
    impacts = db.query(Form3CDImpact).filter(Form3CDImpact.client_id == client_id).all()
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame([_to_dict(s, ["transaction_id", "score", "level", "reasons"]) for s in scores]).to_excel(writer, index=False, sheet_name="Risk Scores")
        pd.DataFrame([_to_dict(a, ["transaction_id", "alert_type", "issue", "severity", "suggested_review"]) for a in alerts]).to_excel(writer, index=False, sheet_name="Statutory Alerts")
        pd.DataFrame([_to_dict(i, ["source_type", "source_id", "clause_area", "observation", "suggested_review"]) for i in impacts]).to_excel(writer, index=False, sheet_name="Form 3CD Impact")
    output.seek(0)
    return output


def gst_reco_excel(db: Session, client_id: int) -> BytesIO:
    return _excel(export_rows(db, client_id), "GST Reco")


def working_paper_docx(db: Session, client_id: int) -> BytesIO:
    client = db.get(Client, client_id)
    paper = db.query(WorkingPaper).filter(WorkingPaper.client_id == client_id).order_by(WorkingPaper.id.desc()).first()
    queries = db.query(ClientQuery).filter(ClientQuery.client_id == client_id).all()
    doc = Document()
    doc.add_heading("AuditXpenser Expense Audit Working Paper", level=1)
    doc.add_paragraph(f"Client name: {client.name if client else client_id}")
    doc.add_paragraph(f"Financial year: {client.financial_year if client else ''}")
    for heading in ["Objective", "Scope", "Data uploaded and reviewed", "Expense areas analysed", "Procedures performed", "Bill matching summary", "Key exceptions", "TDS/GST/RCM observations", "Form 3CD potential impact", "Client queries generated", "CA review notes", "Conclusion placeholder", "Prepared by", "Reviewed by"]:
        doc.add_heading(heading, level=2)
        if paper and heading in {"Objective", "Scope", "Procedures performed", "Conclusion placeholder"}:
            doc.add_paragraph(paper.content)
        elif heading == "Bill matching summary":
            _add_bill_matching_summary(doc, db, client_id)
        elif heading == "Form 3CD potential impact":
            _add_module_review_summary(doc, db, client_id)
        elif heading == "Client queries generated":
            for query in queries:
                doc.add_paragraph(f"{query.query_number}: {query.suggested_wording}", style="List Bullet")
        else:
            doc.add_paragraph("CA Review Required.")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _add_bill_matching_summary(doc: Document, db: Session, client_id: int) -> None:
    rows = db.query(BillMatchingResult).filter(BillMatchingResult.client_id == client_id).all()
    if not rows:
        doc.add_paragraph("No bill matching run data available. CA Review Required.")
        return
    matched = sum(1 for row in rows if row.match_status == "matched")
    high = sum(1 for row in rows if row.risk_level == "High")
    doc.add_paragraph(f"Bill matching results reviewed: {len(rows)}. Matched: {matched}. High-risk / unmatched or mismatched items: {high}.")
    for row in rows[:20]:
        if row.risk_level in {"High", "Medium"}:
            doc.add_paragraph(f"{row.match_status}: {row.bill_vendor_name or row.book_vendor_name or 'Vendor not captured'} | {row.bill_invoice_number or row.book_invoice_number or '-'} | {row.suggested_action or 'CA review required.'}", style="List Bullet")


def _add_module_review_summary(doc: Document, db: Session, client_id: int) -> None:
    depreciation = db.query(FixedAssetDepreciation).filter(FixedAssetDepreciation.client_id == client_id).all()
    alerts = db.query(FixedAssetReviewAlert).filter(FixedAssetReviewAlert.client_id == client_id).all()
    doc.add_paragraph(f"Fixed asset depreciation summary: current year depreciation Rs. {sum(row.depreciation_for_year or 0 for row in depreciation):,.2f}; closing WDV Rs. {sum(row.closing_wdv or 0 for row in depreciation):,.2f}.")
    doc.add_paragraph(f"Fixed asset review alerts: {len(alerts)}. Language is indicative and subject to CA verification.")
    bill_reviews = db.query(BillMatchingResult).filter(BillMatchingResult.client_id == client_id, BillMatchingResult.match_status.in_(["only_in_bill", "only_in_books", "amount_mismatch", "gst_mismatch", "capital_review"])).count()
    doc.add_paragraph(f"Bill matching potential Form 3CD / expense evidence review items: {bill_reviews}. Possible impact only; final reporting requires CA verification.")


def _excel(rows: list[dict], sheet_name: str) -> BytesIO:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame(rows).to_excel(writer, index=False, sheet_name=sheet_name)
    output.seek(0)
    return output


def _to_dict(obj, fields: list[str]) -> dict:
    return {field: getattr(obj, field) for field in fields}
