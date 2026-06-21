from collections import Counter, defaultdict
from datetime import datetime
from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from sqlalchemy.orm import Session

from app.models import AuditException, Client, ClientQuery, UploadedFile
from app.services.exception_register_service import latest_audit_run, map_exception_to_documents_required, map_exception_to_form_3cd_clause


BANNED_WORDS = ["dis" + "allowed", "non-" + "compliant", "default " + "confirmed", "vio" + "lation", "wrong " + "treatment"]
BLUE = "1F4E79"
LIGHT_BLUE = "E8F0FE"
LIGHT_GREY = "F8F9FA"


def get_exception_register_data(db: Session, client_id: int) -> dict:
    client = db.get(Client, client_id)
    run = latest_audit_run(db, client_id)
    query = db.query(AuditException).filter(AuditException.client_id == client_id)
    if run:
        query = query.filter(AuditException.audit_run_id == run.id)
    exceptions = query.order_by(AuditException.risk_level.asc(), AuditException.id.asc()).all()
    category = Counter(item.exception_type for item in exceptions)
    risks = Counter(item.risk_level for item in exceptions)
    clauses = Counter(item.form_3cd_clause or "CA Review Required" for item in exceptions)
    amount_by_category = defaultdict(float)
    for item in exceptions:
        amount_by_category[item.exception_type] += abs(item.amount or 0)
    pending_queries = db.query(ClientQuery).filter(ClientQuery.client_id == client_id, ClientQuery.status == "Pending").count()
    return {
        "client": _client_dict(client),
        "audit_run": _run_dict(run),
        "total_exceptions": len(exceptions),
        "indicative_amount": sum(abs(item.amount or 0) for item in exceptions),
        "category_summary": [
            {"category": key, "count": category[key], "indicative_amount": amount_by_category[key], "risk_level": _category_risk(exceptions, key)}
            for key in sorted(category)
        ],
        "risk_summary": [{"risk_level": key, "count": risks[key]} for key in ["High", "Medium", "Low"] if risks[key]],
        "form_3cd_summary": [{"clause": key, "count": clauses[key]} for key in sorted(clauses)],
        "pending_query_count": pending_queries,
        "exceptions": [_exception_dict(item) for item in exceptions],
    }


def generate_exception_register_xlsx(db: Session, client_id: int) -> BytesIO:
    data = get_exception_register_data(db, client_id)
    wb = Workbook()
    ws = wb.active
    ws.title = "Cover"
    _cover_sheet(ws, data)
    _summary_sheet(wb.create_sheet("Exception Summary"), data)
    _exceptions_sheet(wb.create_sheet("All Exceptions"), data["exceptions"])
    _exceptions_sheet(wb.create_sheet("TDS Exceptions"), _filter(data["exceptions"], ["TDS"]), compact=True)
    _exceptions_sheet(wb.create_sheet("GST Exceptions"), _filter(data["exceptions"], ["GST", "RCM"]), compact=True)
    _exceptions_sheet(wb.create_sheet("Cash & Missing Bill"), _filter(data["exceptions"], ["Cash", "Missing", "Supporting"]), compact=True)
    _exceptions_sheet(wb.create_sheet("Capital vs Revenue"), _filter(data["exceptions"], ["Capital"]), compact=True)
    for sheet in wb.worksheets:
        sheet.sheet_view.showGridLines = True
        _autosize(sheet)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def generate_working_paper_docx(db: Session, client_id: int) -> BytesIO:
    data = get_exception_register_data(db, client_id)
    client = data["client"]
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
        section.header.paragraphs[0].text = f"AuditXpenser | {client['name']} | FY {client['financial_year']} | CONFIDENTIAL"
        section.footer.paragraphs[0].text = "Indicative findings only. CA Review Required."
    _title(doc, "AUDITXPENSER")
    _center(doc, "AI-Assisted Audit Working Paper")
    _title(doc, "TAX AUDIT WORKING PAPER", size=18)
    _center(doc, "Under Section 44AB of the Income Tax Act, 1961")
    _client_table(doc, client)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%d-%b-%Y %H:%M')}")
    doc.add_page_break()
    _heading(doc, "1. Disclaimer and Scope")
    doc.add_paragraph("This working paper has been prepared using AuditXpenser, an AI-assisted audit analysis tool. The analysis is based on uploaded books and supporting records. All findings are indicative in nature and require CA professional review before any audit conclusion is drawn.")
    _heading(doc, "2. Audit Run Summary")
    _key_values(doc, [
        ("Risk Score", f"{data['audit_run']['risk_score']} / 100"),
        ("Risk Label", data["audit_run"]["risk_label"]),
        ("Total Vouchers", data["audit_run"]["total_vouchers"]),
        ("Total Exceptions", data["total_exceptions"]),
        ("Indicative Amount", format_inr(data["indicative_amount"])),
    ])
    _heading(doc, "3. Scope and Basis of Review")
    doc.add_paragraph("The review covers uploaded day book entries, bills, GST data, TDS data and system-generated audit outputs available in AuditXpenser.")
    _heading(doc, "4. Data Sources Used")
    _files_table(doc, db, client_id)
    _heading(doc, "5. Exception Summary")
    _doc_table(doc, ["Category", "Count", "Indicative Amount", "Risk Level"], [[r["category"], r["count"], format_inr(r["indicative_amount"]), r["risk_level"]] for r in data["category_summary"]])
    _heading(doc, "6. Clause-wise Potential Form 3CD Impact")
    _doc_table(doc, ["Clause", "Possible Items"], [[r["clause"], r["count"]] for r in data["form_3cd_summary"]])
    _heading(doc, "7. Detailed Exception Register")
    rows = [[e["voucher_date"] or "", e["voucher_number"] or "", e["party_name"] or "", format_inr(e["amount"] or 0), e["exception_type"], e["risk_level"]] for e in data["exceptions"][:25]]
    _doc_table(doc, ["Date", "Voucher", "Party", "Amount", "Review Area", "Risk"], rows)
    doc.add_paragraph("Full list available in the Exception Register Excel export.")
    _heading(doc, "8. CA Review Notes")
    doc.add_paragraph("CA Review Required. Notes may be added against each exception in the Exceptions tab.")
    _heading(doc, "9. Management Query Summary")
    queries = db.query(ClientQuery).filter(ClientQuery.client_id == client_id).order_by(ClientQuery.id.asc()).limit(20).all()
    for query in queries:
        doc.add_paragraph(f"{query.query_number}: {query.suggested_wording}", style="List Bullet")
    _heading(doc, "10. Disclaimer")
    doc.add_paragraph("This document contains indicative findings only and does not constitute final audit reporting or legal conclusion. The CA must apply independent professional judgement.")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_query_letter_docx(db: Session, client_id: int, status: str = "Pending") -> BytesIO:
    client = db.get(Client, client_id)
    queries = db.query(ClientQuery).filter(ClientQuery.client_id == client_id)
    if status:
        queries = queries.filter(ClientQuery.status == status)
    rows = queries.order_by(ClientQuery.id.asc()).all()
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1.25)
        section.footer.paragraphs[0].text = "Confidential - For CA and Client use only | AuditXpenser generated"
    para = doc.add_paragraph("[CA Firm Name]\n[Address line 1]\n[City, PIN]\n")
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph(f"To,\nThe Management,\n{client.name if client else client_id},\nPAN: {client.pan if client else ''} | GSTIN: {client.gstin if client else ''}")
    subject = doc.add_paragraph("Sub: Tax Audit Queries for Financial Year 2025-26 - U/s 44AB of the Income Tax Act, 1961")
    subject.runs[0].bold = True
    subject.runs[0].underline = True
    doc.add_paragraph("Dear Sir / Madam,")
    doc.add_paragraph("As part of the Tax Audit proceedings, the following queries / clarifications require your written response with supporting documents wherever applicable. A prompt and complete response will assist in timely completion of the audit.")
    _doc_table(doc, ["Query No.", "Voucher Date", "Voucher No.", "Party", "Amount", "Category", "Query / Clarification", "Documents Required", "Client Response", "CA Remarks"], [
        [
            q.query_number,
            q.transaction_date.strftime("%d-%b-%Y") if q.transaction_date else "",
            _voucher_for_query(db, q),
            q.vendor or "",
            format_inr(q.amount or 0),
            q.category or q.issue_detected[:40],
            q.suggested_wording,
            q.documents_required or q.required_document or map_exception_to_documents_required(q.category or ""),
            q.client_response or "",
            q.ca_note or "",
        ]
        for q in rows
    ])
    doc.add_paragraph(f"Total queries raised: {len(rows)} | Total indicative amount under query: {format_inr(sum(q.amount or 0 for q in rows))}")
    doc.add_paragraph("Kindly treat this matter with priority. Incomplete response may delay finalisation of the Tax Audit Report and filing of Form 3CD.")
    doc.add_paragraph("\n\nFor [CA Firm Name]\nChartered Accountants\n\n(CA Name)\nMembership No.: [MCA No.]\nPartner / Proprietor")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def format_inr(amount) -> str:
    return f"₹{indian_number_format(amount)}"


def indian_number_format(amount) -> str:
    try:
        value = float(amount or 0)
    except (TypeError, ValueError):
        value = 0
    sign = "-" if value < 0 else ""
    whole, decimal = f"{abs(value):.2f}".split(".")
    if len(whole) > 3:
        last = whole[-3:]
        rest = whole[:-3]
        groups = []
        while len(rest) > 2:
            groups.insert(0, rest[-2:])
            rest = rest[:-2]
        if rest:
            groups.insert(0, rest)
        whole = ",".join(groups + [last])
    return f"{sign}{whole}.{decimal}"


def safe_text(value) -> str:
    return "" if value is None else str(value)


def scan_banned_words(text: str) -> list[str]:
    lower = text.lower()
    return [word for word in BANNED_WORDS if word in lower]


def _cover_sheet(ws, data):
    ws.merge_cells("A1:H1")
    ws["A1"] = "AUDITXPENSER - EXCEPTION REGISTER"
    ws["A1"].font = Font(bold=True, size=16, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor=BLUE)
    client = data["client"]
    rows = [
        ("Client Name:", client["name"]),
        ("PAN:", client["pan"]),
        ("GSTIN:", client["gstin"]),
        ("Financial Year:", client["financial_year"]),
        ("Report Generated:", datetime.now().strftime("%d-%b-%Y %H:%M")),
        ("Prepared by:", "AuditXpenser (AI-assisted audit tool)"),
        ("DISCLAIMER", "This report is based on rule-based analysis of uploaded records. All exceptions are indicative and require CA professional review before any audit conclusion is drawn."),
        ("Indicative Risk Score:", data["audit_run"]["risk_score"]),
        ("Total Exceptions Flagged:", data["total_exceptions"]),
        ("Indicative Amount Under Review:", format_inr(data["indicative_amount"])),
    ]
    for idx, row in enumerate(rows, start=3):
        ws.cell(idx, 1, row[0]).font = Font(bold=True)
        ws.cell(idx, 2, row[1])


def _summary_sheet(ws, data):
    headers = ["Sr. No.", "Exception Category", "No. of Exceptions", "Indicative Amount", "Risk Level", "Potential Form 3CD Clause", "Remarks"]
    _write_header(ws, headers)
    for idx, row in enumerate(data["category_summary"], start=2):
        ws.append([idx - 1, row["category"], row["count"], row["indicative_amount"], row["risk_level"], map_exception_to_form_3cd_clause(row["category"]), "CA Review Required"])
    ws.append(["", "TOTAL", data["total_exceptions"], data["indicative_amount"], "", "", "All figures are indicative."])
    _style_table(ws)


def _exceptions_sheet(ws, rows, compact=False):
    headers = ["Sr. No.", "Voucher Date", "Voucher Type", "Voucher No.", "Party Name", "Ledger Name", "Amount", "Exception Type", "Exception Description", "Risk Level", "Form 3CD Clause", "CA Remarks"]
    _write_header(ws, headers)
    for idx, item in enumerate(rows, start=1):
        ws.append([idx, item["voucher_date"], item["voucher_type"], item["voucher_number"], item["party_name"], item["ledger_name"], item["amount"], item["exception_type"], item["exception_description"], item["risk_level"], item["form_3cd_clause"], item["ca_remarks"] or ""])
    _style_table(ws)


def _write_header(ws, headers):
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.alignment = Alignment(horizontal="center")
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def _style_table(ws):
    for row in ws.iter_rows(min_row=2):
        fill = PatternFill("solid", fgColor=LIGHT_GREY if row[0].row % 2 == 0 else "FFFFFF")
        for cell in row:
            cell.fill = fill
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if isinstance(cell.value, (int, float)) and cell.column >= 3:
                cell.number_format = '₹#,##0.00'


def _autosize(ws):
    for column in ws.columns:
        width = max(len(safe_text(cell.value)) for cell in column)
        ws.column_dimensions[get_column_letter(column[0].column)].width = min(max(width + 2, 10), 45)


def _filter(rows, tokens):
    return [row for row in rows if any(token.lower() in (row["exception_type"] or "").lower() for token in tokens)]


def _client_dict(client):
    return {"id": client.id if client else None, "name": client.name if client else "", "pan": client.pan if client else "", "gstin": client.gstin if client else "", "financial_year": client.financial_year if client else "2025-26"}


def _run_dict(run):
    if not run:
        return {"id": None, "run_at": None, "risk_score": 0, "risk_label": "Low", "total_vouchers": 0, "total_exceptions": 0, "indicative_amount": 0}
    return {"id": run.id, "run_at": run.run_at, "risk_score": run.risk_score, "risk_label": run.risk_label, "total_vouchers": run.total_vouchers, "total_exceptions": run.total_exceptions, "indicative_amount": run.indicative_amount}


def _exception_dict(item):
    return {
        "id": item.id,
        "audit_run_id": item.audit_run_id,
        "transaction_id": item.transaction_id,
        "voucher_date": item.voucher_date.isoformat() if item.voucher_date else None,
        "voucher_type": item.voucher_type,
        "voucher_number": item.voucher_number,
        "party_name": item.party_name,
        "ledger_name": item.ledger_name,
        "amount": item.amount,
        "exception_type": item.exception_type,
        "exception_description": item.exception_description,
        "risk_level": item.risk_level,
        "form_3cd_clause": item.form_3cd_clause,
        "status": item.status,
        "ca_remarks": item.ca_remarks,
    }


def _category_risk(exceptions, category):
    values = [item.risk_level for item in exceptions if item.exception_type == category]
    if "High" in values:
        return "High"
    if "Medium" in values:
        return "Medium"
    return "Low"


def _title(doc, text, size=16):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def _center(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _heading(doc, text):
    paragraph = doc.add_heading(text, level=1)
    paragraph.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def _client_table(doc, client):
    rows = [("Client Name", client["name"]), ("PAN", client["pan"]), ("GSTIN", client["gstin"]), ("Financial Year", client["financial_year"]), ("Nature of Audit", "Tax Audit u/s 44AB")]
    _doc_table(doc, ["Particular", "Details"], rows)


def _key_values(doc, rows):
    _doc_table(doc, ["Metric", "Value"], rows)


def _files_table(doc, db, client_id):
    files = db.query(UploadedFile).filter(UploadedFile.client_id == client_id).order_by(UploadedFile.category.asc()).all()
    _doc_table(doc, ["File Type", "File Name", "Status"], [[f.category, f.filename, f"{f.parse_status} - {f.records_extracted} records"] for f in files[:30]])


def _doc_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = safe_text(header)
        cell.paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = safe_text(value)
    return table


def _voucher_for_query(db, query):
    if not query.exception_id:
        return ""
    exception = db.get(AuditException, query.exception_id)
    return exception.voucher_number if exception else ""
